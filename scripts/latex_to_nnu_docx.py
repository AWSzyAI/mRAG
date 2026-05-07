#!/usr/bin/env python3
"""Convert the local thesis LaTeX source into an editable NNU-style DOCX."""

from __future__ import annotations

import argparse
import shutil
import tempfile
import re
from copy import deepcopy
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from lxml import etree
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


CHINESE_FONT = "SimSun"
WESTERN_FONT = "Times New Roman"
HEADER_TEXT = "南京师范大学计算机与电子信息学院本科毕业论文"
REF_MAP: dict[str, str] = {}
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
ASSET_DIR: Path | None = None
WORD_IMAGE_DIR: Path | None = None
SONGTI_FONT = "/System/Library/Fonts/Supplemental/Songti.ttc"
TIMES_FONT = "/System/Library/Fonts/Supplemental/Times New Roman.ttf"


def set_run_font(run, size_pt: float | None = None, bold: bool | None = None, italic: bool | None = None):
    run.font.name = WESTERN_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), WESTERN_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), WESTERN_FONT)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_paragraph_format(paragraph, first_line=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Cm(0.74) if first_line else None
    paragraph.alignment = align


def set_style_font(style, size_pt: float, bold=False, italic=False, font_name: str | None = None):
    font_name = font_name or WESTERN_FONT
    font = style.font
    font.name = font_name
    font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    font.color.rgb = RGBColor(0, 0, 0)
    if style._element.rPr is None:
        style._element.append(OxmlElement("w:rPr"))
    r_fonts = style._element.rPr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        style._element.rPr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), CHINESE_FONT)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)


def get_or_create_paragraph_style(styles, name: str):
    try:
        return styles[name]
    except KeyError:
        return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def clear_style_color(style):
    if style.element.rPr is None:
        return
    for color in style.element.rPr.findall(qn("w:color")):
        style.element.rPr.remove(color)


def add_page_number(paragraph, fmt: str | None = None):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_text, fld_end])
    set_run_font(run, 9)


def set_page_number_type(section, start: int, fmt: str):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))
    pg_num.set(qn("w:fmt"), fmt)


def set_header_footer(section, enabled=True, page_fmt="decimal", start=1):
    section.top_margin = Cm(2.8)
    section.bottom_margin = Cm(2.8)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)
    section.header_distance = Cm(1.4)
    section.footer_distance = Cm(1.4)
    set_page_number_type(section, start, page_fmt)
    if not enabled:
        return
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = HEADER_TEXT
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(hp.runs[0], 9)
    p_pr = hp._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    border.append(bottom)
    p_pr.append(border)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    add_page_number(fp)


def add_field(paragraph, field_code: str, placeholder: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, text, end])
    set_run_font(run, 12)


def add_toc(document: Document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目  录")
    set_run_font(r, 16, bold=True)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    toc = document.add_paragraph()
    add_field(toc, r'TOC \o "1-3" \h \z \u', "右键更新域以生成目录")
    set_paragraph_format(toc, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)


def get_macro(text: str, name: str, default=""):
    m = re.search(rf"\\newcommand\{{\\{name}\}}\{{(.*?)\}}", text)
    return latex_to_text(m.group(1)) if m else default


SUBSCRIPT_CHARS = {
    **dict(zip("0123456789", "₀₁₂₃₄₅₆₇₈₉")),
    "a": "ₐ", "e": "ₑ", "h": "ₕ", "i": "ᵢ", "j": "ⱼ", "k": "ₖ", "l": "ₗ",
    "m": "ₘ", "n": "ₙ", "o": "ₒ", "p": "ₚ", "r": "ᵣ", "s": "ₛ", "t": "ₜ",
    "u": "ᵤ", "v": "ᵥ", "x": "ₓ", "+": "₊", "-": "₋", "=": "₌", "(": "₍", ")": "₎",
}


def to_subscript(text: str) -> str:
    if all(ch in SUBSCRIPT_CHARS for ch in text):
        return "".join(SUBSCRIPT_CHARS[ch] for ch in text)
    return "_" + text


def math_to_text(expr: str) -> str:
    expr = expr.strip()
    replacements = {
        r"\mathcal{C}": "𝒞",
        r"\mathcal{D}": "𝒟",
        r"\hat{y}": "ŷ",
        r"\dots": "…",
        r"\ldots": "…",
        r"\to": "→",
        r"\in": "∈",
        r"\sum": "Σ",
        r"\times": "×",
        r"\geq": "≥",
        r"\leq": "≤",
        r"\mathrm": "",
        r"\{": "{",
        r"\}": "}",
    }
    for src, dst in replacements.items():
        expr = expr.replace(src, dst)
    expr = re.sub(r"\\mathcal\{([^{}]+)\}", r"\1", expr)
    expr = re.sub(r"\\hat\{([^{}]+)\}", r"^\1", expr)
    expr = re.sub(r"\\mathrm\{([^{}]+)\}", r"\1", expr)
    expr = re.sub(r"_\{([^{}]+)\}", lambda m: to_subscript(m.group(1)), expr)
    expr = re.sub(r"_([A-Za-z0-9])", lambda m: to_subscript(m.group(1)), expr)
    expr = re.sub(r"\^\{([^{}]+)\}", r"^(\1)", expr)
    expr = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", expr)
    expr = re.sub(r"\\[a-zA-Z]+", "", expr)
    expr = expr.replace("{", "").replace("}", "")
    return re.sub(r"\s+", " ", expr).strip()


def latex_to_text(text: str, citation_map: dict[str, str] | None = None) -> str:
    citation_map = citation_map or {}
    text = text.replace("\n", " ")
    text = re.sub(r"(?<!\\)%.*", "", text)
    text = re.sub(r"\\zihao\{[^{}]*\}", "", text)
    text = re.sub(r"\\hspace\*?\{[^{}]*\}", "", text)
    text = re.sub(r"\\vspace\{[^{}]*\}", "", text)
    text = re.sub(r"\\begin\{[^{}]*\}(?:\[[^\]]*\])?(?:\{[^{}]*\})?", "", text)
    text = re.sub(r"\\end\{[^{}]*\}", "", text)
    text = re.sub(r"\\noindent", "", text)
    text = text.replace("~", " ")
    replacements = {
        r"\%": "%",
        r"\&": "&",
        r"\_": "_",
        r"\#": "#",
        r"\{": "<<LBRACE>>",
        r"\}": "<<RBRACE>>",
        r"\quad": "  ",
        r"\qquad": "    ",
        r"\dots": "...",
        r"\ldots": "...",
        r"\to": "→",
        r"\geq": "≥",
        r"\leq": "≤",
        r"\times": "×",
        r"\mathrm": "",
        r"``": "“",
        r"''": "”",
        r"--": "-",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)

    def cite_repl(match):
        keys = [k.strip() for k in match.group(1).split(",")]
        nums = [citation_map.get(k, k) for k in keys]
        return "<CITE:" + ",".join(nums) + ">"

    text = re.sub(r"\$([^$]+)\$", lambda m: math_to_text(m.group(1)), text)
    text = re.sub(r"\\cite[t|p|alp]*\{([^{}]+)\}", cite_repl, text)
    text = re.sub(r"\\href\{[^{}]*\}\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\url\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\mathcal\{([^{}]+)\}", r"\1", text)
    text = re.sub(r"\\hat\{([^{}]+)\}", r"hat(\1)", text)
    text = re.sub(r"\\mathrm\{([^{}]+)\}", r"\1", text)
    text = re.sub(r"\\textbf\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\textit\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\emph\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\texttt\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\textcolor\{[^{}]*\}\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\modelcap\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\best\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\worst\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\secondbest\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\secondworst\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\focusval\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\upperbase\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\lowerbase\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\ref\{([^{}]*)\}", lambda m: REF_MAP.get(m.group(1), m.group(1)), text)
    text = re.sub(r"\\label\{[^{}]*\}", "", text)
    text = re.sub(r"\\[a-zA-Z]+", "", text)
    text = text.replace("{", "").replace("}", "")
    text = text.replace("<<LBRACE>>", "{").replace("<<RBRACE>>", "}")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def add_text_with_citations(paragraph, text: str, size=12, bold=False):
    pos = 0
    for match in re.finditer(r"<CITE:([^>]+)>", text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size, bold=bold)
        run = paragraph.add_run("[" + match.group(1) + "]")
        set_run_font(run, size, bold=bold)
        run.font.superscript = True
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size, bold=bold)


def max_word_id(root, tag: str, attr_name: str) -> int:
    values = []
    for el in root.findall(qn(tag)):
        value = el.get(qn(attr_name))
        if value and value.isdigit():
            values.append(int(value))
    return max(values or [0])


def ensure_reference_numbering(document) -> str:
    existing = getattr(document, "_reference_num_id", None)
    if existing is not None:
        return existing

    numbering = document.part.numbering_part.element
    abstract_id = str(max_word_id(numbering, "w:abstractNum", "w:abstractNumId") + 1)
    num_id = str(max_word_id(numbering, "w:num", "w:numId") + 1)

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), abstract_id)
    multi_level_type = OxmlElement("w:multiLevelType")
    multi_level_type.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level_type)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "[%1]")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "space")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl_ppr = OxmlElement("w:pPr")
    lvl_ind = OxmlElement("w:ind")
    lvl_ind.set(qn("w:left"), "0")
    lvl_ind.set(qn("w:hanging"), "0")
    lvl_ppr.append(lvl_ind)
    lvl_rpr = OxmlElement("w:rPr")
    lvl_fonts = OxmlElement("w:rFonts")
    lvl_fonts.set(qn("w:ascii"), WESTERN_FONT)
    lvl_fonts.set(qn("w:hAnsi"), WESTERN_FONT)
    lvl_fonts.set(qn("w:eastAsia"), CHINESE_FONT)
    lvl_size = OxmlElement("w:sz")
    lvl_size.set(qn("w:val"), "24")
    lvl_rpr.append(lvl_fonts)
    lvl_rpr.append(lvl_size)
    for child in (start, num_fmt, lvl_text, suffix, lvl_jc, lvl_ppr, lvl_rpr):
        lvl.append(child)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_id)
    num.append(abstract_ref)
    numbering.append(num)

    document._reference_num_id = num_id
    return num_id


def add_reference_paragraph(document, ref: str, num_id: str):
    p = document.add_paragraph()
    set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), num_id)
    numPr.append(ilvl)
    numPr.append(num)
    pPr.append(numPr)
    add_text_with_citations(p, ref, 12)


def add_body_paragraph(document, text: str, citation_map: dict[str, str]):
    text = latex_to_text(text, citation_map)
    if not text:
        return
    p = document.add_paragraph()
    set_paragraph_format(p)
    add_text_with_citations(p, text, 12)


def add_subsubsection_heading(document, title: str):
    title = latex_to_text(title)
    if not title:
        return
    p = document.add_paragraph()
    set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, 12, bold=True)


def add_center_heading(document, title: str, page_break=False):
    if page_break and document.paragraphs:
        document.add_page_break()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title)
    set_run_font(r, 16, bold=True)
    return p


def add_chapter(document, number: int | None, title: str):
    document.add_page_break()
    text = f"第 {number} 章  {title}" if number is not None else title
    p = document.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, 16, bold=True)


def add_section_heading(document, level: int, number: str, title: str):
    style = "Heading 2" if level == 2 else "Heading 3"
    p = document.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = None
    run = p.add_run(f"{number}  {title}")
    set_run_font(run, 14 if level == 2 else 12, bold=True)


def add_cover(document, meta: dict[str, str]):
    for text, size in [("南  京  师  范  大  学", 18), ("毕  业  设  计（论  文）", 18), ("（2026 届）", 14)]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(text)
        set_run_font(r, size, bold=True)
    document.add_paragraph()
    fields = [
        ("题    目", meta.get("title", "")),
        ("学    院", meta.get("college", "")),
        ("专    业", meta.get("major", "")),
        ("姓    名", meta.get("author", "")),
        ("学    号", meta.get("id", "")),
        ("指导教师", meta.get("advisor", "")),
    ]
    for label, value in fields:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(3.0)
        p.paragraph_format.space_after = Pt(14)
        r = p.add_run(f"{label}：  {value}")
        set_run_font(r, 14)
    for _ in range(5):
        document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("南京师范大学教务处  制")
    set_run_font(r, 12)


def split_main_parts(main_text: str):
    abstract_cn = re.search(r"\\begin\{center\}\s*\{\\zihao\{3\}\\bfseries 摘\\quad 要\\par\}\s*\\end\{center\}(.*?)\{\\small\\noindent\\textbf\{关键词：\}(.*?)\\par\}", main_text, re.S)
    abstract_en = re.search(r"\\begin\{center\}\s*\{\\zihao\{3\}\\bfseries Abstract\\par\}\s*\\end\{center\}(.*?)\{\\small\\noindent\\textbf\{Key words:\}(.*?)\\par\}", main_text, re.S)
    cn_text = abstract_cn.group(1) if abstract_cn else ""
    cn_keywords = abstract_cn.group(2) if abstract_cn else ""
    en_text = abstract_en.group(1) if abstract_en else ""
    en_keywords = abstract_en.group(2) if abstract_en else ""
    after_content = main_text.split(r"\input{content.tex}", 1)[1].split(r"\end{document}", 1)[0]
    return cn_text, cn_keywords, en_text, en_keywords, after_content


def clean_abstract_block(text: str) -> str:
    text = "\n".join(line for line in text.splitlines() if not line.strip().startswith(r"\addcontentsline"))
    text = re.sub(r"\\texorpdfstring\{[^{}]*\}\{[^{}]*\}", "", text)
    text = re.sub(r"\\noindent\\hspace\*?\{[^{}]*\}%?", "", text)
    return text.strip()


def add_abstracts(document, cn_text, cn_keywords, en_text, en_keywords, citation_map):
    add_center_heading(document, "摘  要")
    for para in re.split(r"\n\s*\n", clean_abstract_block(cn_text)):
        add_body_paragraph(document, para, citation_map)
    p = document.add_paragraph()
    set_paragraph_format(p, first_line=False)
    r = p.add_run("关键词：")
    set_run_font(r, 12, bold=True)
    add_text_with_citations(p, latex_to_text(cn_keywords, citation_map), 12)
    document.add_page_break()

    add_center_heading(document, "Abstract")
    for para in re.split(r"\n\s*\n", clean_abstract_block(en_text)):
        add_body_paragraph(document, para, citation_map)
    p = document.add_paragraph()
    set_paragraph_format(p, first_line=False)
    r = p.add_run("Key words: ")
    set_run_font(r, 12, bold=True)
    add_text_with_citations(p, latex_to_text(en_keywords, citation_map), 12)


def read_braced_argument(text: str, open_brace_index: int) -> str:
    depth = 0
    chars: list[str] = []
    for idx in range(open_brace_index, len(text)):
        ch = text[idx]
        prev = text[idx - 1] if idx > 0 else ""
        if ch == "{" and prev != "\\":
            depth += 1
            if depth == 1:
                continue
        elif ch == "}" and prev != "\\":
            depth -= 1
            if depth == 0:
                return "".join(chars)
        if depth >= 1:
            chars.append(ch)
    return ""


def extract_caption(block: str) -> str:
    match = re.search(r"\\caption(?:\[[^\]]*\])?\s*\{", block, re.S)
    if not match:
        return ""
    return latex_to_text(read_braced_argument(block, match.end() - 1))


def extract_algorithm_caption(block: str, number: str) -> str:
    match = re.search(r"\\textbf\{算法\s*" + re.escape(number) + r"[：:]\s*", block)
    if not match:
        return "E8 多维检索式 RAG 推理过程"
    title = read_braced_argument(block, match.start() + len("\\textbf"))
    title = latex_to_text(title)
    title = re.sub(rf"^算法\s*{re.escape(number)}[：:]\s*", "", title).strip()
    return title or "E8 多维检索式 RAG 推理过程"


def add_caption(document, prefix: str, caption: str, before=False):
    if not caption:
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6 if before else 0)
    p.paragraph_format.space_after = Pt(0 if before else 6)
    r = p.add_run(f"{prefix}{caption}")
    set_run_font(r, 10.5)


def asset_path(name: str) -> Path:
    if ASSET_DIR is None:
        raise RuntimeError("ASSET_DIR is not initialized")
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    return ASSET_DIR / name


def word_image_path(name: str) -> Path | None:
    if WORD_IMAGE_DIR is None:
        return None
    path = WORD_IMAGE_DIR / name
    return path if path.exists() else None


def add_screenshot(document, image: Path, width_cm: float = 14.5):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(image), width=Cm(width_cm))


def render_formula_image(formula: str, state: dict) -> Path | None:
    clean = re.sub(r"\\label\{[^{}]*\}", "", formula).strip()
    if not clean:
        return None
    state["formula_image"] = state.get("formula_image", 0) + 1
    out = asset_path(f"formula_{state['formula_image']:03d}.png")
    try:
        aligned = re.search(r"\\begin\{aligned\}(.*?)\\end\{aligned\}", clean, re.S)
        if aligned:
            lines = []
            for line in re.split(r"\\\\", aligned.group(1)):
                line = line.replace("&", "").strip()
                if line:
                    lines.append(line)
            fig = plt.figure(figsize=(9.0, max(1.15, 0.48 * len(lines))), dpi=240)
            fig.patch.set_alpha(0)
            ax = fig.add_axes([0, 0, 1, 1])
            ax.axis("off")
            y0 = 0.5 + (len(lines) - 1) * 0.16
            for idx, line in enumerate(lines):
                ax.text(0.5, y0 - idx * 0.32, f"${line}$", ha="center", va="center", fontsize=25, color="black")
            fig.savefig(out, transparent=True, bbox_inches="tight", pad_inches=0.03)
            plt.close(fig)
            return out

        fig = plt.figure(figsize=(9.0, 1.15), dpi=240)
        fig.patch.set_alpha(0)
        ax = fig.add_axes([0, 0, 1, 1])
        ax.axis("off")
        ax.text(0.5, 0.5, f"${clean}$", ha="center", va="center", fontsize=26, color="black")
        fig.savefig(out, transparent=True, bbox_inches="tight", pad_inches=0.03)
        plt.close(fig)
        return out
    except Exception:
        plt.close("all")
        return None


def latex_block_to_lines(block: str) -> list[str]:
    text = block
    text = re.sub(r"\\begin\{center\}|\\end\{center\}", "", text)
    text = re.sub(r"\\begin\{minipage\}\{[^{}]*\}|\\end\{minipage\}", "", text)
    text = re.sub(r"\\fbox\{%?", "", text)
    text = re.sub(r"\\setlength\{[^{}]*\}\{[^{}]*\}", "", text)
    text = re.sub(r"\\small|\\hrule|\\par|\\noindent", "", text)
    text = re.sub(r"\\vspace\{[^{}]*\}", "\n", text)
    text = text.replace("\\\\", "\n")
    text = re.sub(r"\\\[(.*?)\\\]", lambda m: "\n" + math_to_text(m.group(1)) + "\n", text, flags=re.S)
    text = latex_to_text(text)
    text = text.replace("Step ", "\nStep ")
    text = text.replace("输入：", "\n输入：").replace("输出：", "\n输出：").replace("参数：", "\n参数：").replace("运行设置：", "\n运行设置：")
    lines = [line.strip(" }%") for line in text.splitlines()]
    return [line for line in lines if line]


def wrap_line(draw, text: str, font, max_width: int) -> list[str]:
    if draw.textlength(text, font=font) <= max_width:
        return [text]
    result: list[str] = []
    current = ""
    for ch in text:
        candidate = current + ch
        if current and draw.textlength(candidate, font=font) > max_width:
            result.append(current)
            current = ch
        else:
            current = candidate
    if current:
        result.append(current)
    return result


def render_pseudocode_image(block: str, state: dict) -> Path | None:
    lines = latex_block_to_lines(block)
    if not lines:
        return None
    state["pseudo_image"] = state.get("pseudo_image", 0) + 1
    out = asset_path(f"pseudocode_{state['pseudo_image']:03d}.png")
    font = ImageFont.truetype(SONGTI_FONT, 28)
    bold_font = ImageFont.truetype(SONGTI_FONT, 30)
    width = 1800
    margin = 70
    line_gap = 16
    temp = Image.new("RGB", (width, 100), "white")
    draw = ImageDraw.Draw(temp)
    wrapped: list[tuple[str, bool]] = []
    for line in lines:
        is_title = line.startswith("算法")
        max_width = width - 2 * margin
        for idx, part in enumerate(wrap_line(draw, line, bold_font if is_title else font, max_width)):
            wrapped.append((part, is_title and idx == 0))
    line_height = 40
    height = 2 * margin + len(wrapped) * (line_height + line_gap) + 20
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle([25, 25, width - 25, height - 25], outline="black", width=2)
    y = margin
    for text, is_title in wrapped:
        current_font = bold_font if is_title else font
        x = (width - draw.textlength(text, font=current_font)) / 2 if is_title else margin
        draw.text((x, y), text, fill="black", font=current_font)
        y += line_height + line_gap
    image.save(out)
    return out


def add_center_block(document, block: str, state: dict, citation_map: dict[str, str]):
    if "minipage" in block or "算法" in block or "Step" in block:
        algorithm_number = f"{state['chapter']}-1"
        screenshot = word_image_path(f"算法{algorithm_number}.png")
        if screenshot:
            add_caption(document, f"算法 {algorithm_number} ", extract_algorithm_caption(block, algorithm_number), before=True)
            add_screenshot(document, screenshot)
            return
        image = render_pseudocode_image(block, state)
        if image and image.exists():
            add_caption(document, f"算法 {algorithm_number} ", extract_algorithm_caption(block, algorithm_number), before=True)
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(image), width=Cm(14.5))
            return
    add_body_paragraph(document, block, citation_map)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, edge_data in {"top": top, "bottom": bottom, "left": left, "right": right}.items():
        if edge_data is None:
            continue
        tag = "w:" + edge_name
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in edge_data.items():
            element.set(qn("w:" + key), str(value))


def apply_three_line_table_borders(table):
    no_border = {"val": "nil", "sz": "0", "color": "auto"}
    line = {"val": "single", "sz": "8", "color": "000000", "space": "0"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=no_border, bottom=no_border, left=no_border, right=no_border)
    if not table.rows:
        return
    for cell in table.rows[0].cells:
        set_cell_border(cell, top=line, bottom=line)
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom=line)


def add_figure(document, block: str, image_dir: Path, state: dict):
    state["figure"] += 1
    imgs = re.findall(r"\\includegraphics(?:\[[^\]]*\])?\{([^{}]+)\}", block)
    for img in imgs:
        path = image_dir / img
        if path.exists():
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(path), width=Cm(14.5))
    add_caption(document, f"图 {state['chapter']}-{state['figure']} ", extract_caption(block), before=False)


def parse_tabular(block: str, citation_map: dict[str, str]):
    m = re.search(r"\\begin\{tabular\}\{.*?\}(.*?)\\end\{tabular\}", block, re.S)
    if not m:
        return []
    body = m.group(1)
    body = re.sub(r"\\(toprule|midrule|bottomrule)", "", body)
    body = re.sub(r"\\cmidrule(?:\([^)]*\))?\{[^{}]*\}", "", body)
    rows = []
    for raw in re.split(r"\\\\", body):
        raw = raw.strip()
        if not raw:
            continue
        cells = [latex_to_text(c.strip(), citation_map) for c in raw.split("&")]
        if any(cells):
            rows.append(cells)
    return rows


def add_table(document, block: str, citation_map: dict[str, str], state: dict):
    state["table"] += 1
    table_number = f"{state['chapter']}-{state['table']}"
    screenshot = word_image_path(f"表{table_number}.png")
    caption = extract_caption(block)
    if screenshot:
        add_caption(document, f"表 {table_number} ", caption, before=True)
        add_screenshot(document, screenshot)
        return
    add_caption(document, f"表 {table_number} ", caption, before=True)
    rows = parse_tabular(block, citation_map)
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = document.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = row[j] if j < len(row) else ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            set_run_font(r, 10.5, bold=(i == 0))
    apply_three_line_table_borders(table)


def add_equation(document, block: str, state: dict):
    state["equation"] += 1
    body = re.sub(r"\\begin\{equation\}|\\end\{equation\}|\\\[|\\\]", "", block).strip()
    image = render_formula_image(body, state)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    if image and image.exists():
        p.add_run().add_picture(str(image), width=Cm(13.5))
    else:
        r = p.add_run(math_to_text(body))
        set_run_font(r, 12)


def add_enumerate(document, block: str, citation_map: dict[str, str]):
    items = re.split(r"\\item", block)
    for idx, item in enumerate(items[1:], 1):
        item = re.sub(r"\\end\{enumerate\}", "", item)
        text = latex_to_text(item, citation_map)
        if not text:
            continue
        p = document.add_paragraph()
        set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.37)
        add_text_with_citations(p, f"{idx}. {text}", 12)


def add_research_items(document, block: str, citation_map: dict[str, str]):
    items = re.split(r"\\item", block)
    for idx, item in enumerate(items[1:], 1):
        item = re.sub(r"\\end\{researchitems\}", "", item)
        text = latex_to_text(item, citation_map)
        if not text:
            continue
        p = document.add_paragraph()
        set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.74)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(17.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        add_text_with_citations(p, f"[{idx}] {text}", 12)


def build_citation_map(bbl_path: Path) -> dict[str, str]:
    if not bbl_path.exists():
        return {}
    text = bbl_path.read_text(encoding="utf-8", errors="ignore")
    keys = re.findall(r"\\bibitem(?:\[[^\]]*\])?\{([^{}]+)\}", text)
    return {key: str(i + 1) for i, key in enumerate(keys)}


def build_ref_map(text: str) -> dict[str, str]:
    ref_map: dict[str, str] = {}
    chapter = section = subsection = figure = table = equation = 0
    pending_number: str | None = None
    pending_kind: str | None = None
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        chapter_match = re.match(r"\\chapter\{(.+?)\}", line)
        if chapter_match:
            chapter += 1
            section = subsection = figure = table = equation = 0
            pending_number = str(chapter)
            pending_kind = "chapter"
            i += 1
            continue
        section_match = re.match(r"\\section\{(.+?)\}", line)
        if section_match:
            section += 1
            subsection = 0
            pending_number = f"{chapter}.{section}"
            pending_kind = "section"
            i += 1
            continue
        subsection_match = re.match(r"\\subsection\{(.+?)\}", line)
        if subsection_match:
            subsection += 1
            pending_number = f"{chapter}.{section}.{subsection}"
            pending_kind = "subsection"
            i += 1
            continue

        env_match = re.match(r"\\begin\{(figure|table|equation)\}", line)
        if env_match:
            env = env_match.group(1)
            block = [line]
            while i + 1 < len(lines):
                i += 1
                block.append(lines[i])
                if re.search(rf"\\end\{{{env}\}}", lines[i]):
                    break
            block_text = "\n".join(block)
            if env == "figure":
                figure += 1
                number = f"{chapter}-{figure}"
            elif env == "table":
                table += 1
                number = f"{chapter}-{table}"
            else:
                equation += 1
                number = f"{chapter}-{equation}"
            for label in re.findall(r"\\label\{([^{}]+)\}", block_text):
                ref_map[label] = number
            i += 1
            continue

        label_match = re.match(r"\\label\{([^{}]+)\}", line)
        if label_match and pending_number:
            ref_map[label_match.group(1)] = pending_number
            pending_kind = None
            i += 1
            continue

        if line and not line.startswith("\\label") and pending_kind not in {"chapter", "section", "subsection"}:
            pending_number = None
        i += 1
    return ref_map


def read_bibliography(bbl_path: Path) -> list[str]:
    if not bbl_path.exists():
        return []
    text = bbl_path.read_text(encoding="utf-8", errors="ignore")
    entries = re.split(r"\\bibitem(?:\[[^\]]*\])?\{[^{}]+\}", text)[1:]
    refs = []
    for entry in entries:
        entry = re.sub(r"\\newblock", " ", entry)
        entry = re.sub(r"\\begin\{thebibliography\}\{[^{}]*\}|\\end\{thebibliography\}", "", entry)
        clean = latex_to_text(entry)
        if clean:
            refs.append(clean)
    return refs


def parse_blocks(text: str):
    pattern = re.compile(
        r"\\begin\{(?:figure|table|equation|enumerate|researchitems|center|verbatim)\}.*?\\end\{(?:figure|table|equation|enumerate|researchitems|center|verbatim)\}|\\\[.*?\\\]",
        re.S,
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            yield "text", text[pos:m.start()]
        block = m.group(0)
        kind_match = re.match(r"\\begin\{([^{}]+)\}", block)
        kind = kind_match.group(1) if kind_match else "equation"
        yield kind, block
        pos = m.end()
    if pos < len(text):
        yield "text", text[pos:]


def process_text_chunk(document, chunk: str, state: dict, citation_map: dict[str, str]):
    para_lines = []

    def flush():
        if para_lines:
            add_body_paragraph(document, " ".join(para_lines), citation_map)
            para_lines.clear()

    for line in chunk.splitlines():
        stripped = line.strip()
        if not stripped or stripped.startswith("%"):
            flush()
            continue
        if stripped.startswith(("\\label", "\\centering", "\\resizebox", "\\toprule", "\\midrule", "\\bottomrule", "\\graphicspath")):
            continue
        if stripped.startswith("\\bibliographystyle"):
            continue
        if stripped.startswith("\\bibliography"):
            flush()
            refs = state.get("references", [])
            add_chapter(document, None, "参考文献")
            ref_num_id = ensure_reference_numbering(document)
            for ref in refs:
                add_reference_paragraph(document, ref, ref_num_id)
            continue
        m = re.match(r"\\chapter\{(.+?)\}", stripped)
        if m:
            flush()
            state["chapter"] += 1
            state["section"] = 0
            state["subsection"] = 0
            state["figure"] = 0
            state["table"] = 0
            state["equation"] = 0
            add_chapter(document, state["chapter"], latex_to_text(m.group(1)))
            continue
        m = re.match(r"\\chapter\*\{(.+?)\}", stripped)
        if m:
            flush()
            add_chapter(document, None, latex_to_text(m.group(1)))
            continue
        m = re.match(r"\\section\{(.+?)\}", stripped)
        if m:
            flush()
            state["section"] += 1
            state["subsection"] = 0
            add_section_heading(document, 2, f"{state['chapter']}.{state['section']}", latex_to_text(m.group(1)))
            continue
        m = re.match(r"\\section\*\{(.+?)\}", stripped)
        if m:
            flush()
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(latex_to_text(m.group(1)))
            set_run_font(r, 14, bold=True)
            continue
        m = re.match(r"\\researchheading\{(.+?)\}", stripped)
        if m:
            flush()
            p = document.add_paragraph()
            set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(latex_to_text(m.group(1)) + "：")
            set_run_font(r, 12, bold=True)
            continue
        m = re.match(r"\\subsection\{(.+?)\}", stripped)
        if m:
            flush()
            state["subsection"] += 1
            state["subsubsection"] = 0
            add_section_heading(document, 3, f"{state['chapter']}.{state['section']}.{state['subsection']}", latex_to_text(m.group(1)))
            continue
        m = re.match(r"\\subsubsection\{(.+?)\}", stripped)
        if m:
            flush()
            state["subsubsection"] += 1
            add_subsubsection_heading(document, m.group(1))
            continue
        if stripped.startswith(("\\addcontentsline", "\\clearpage", "\\cleardoublepage", "\\par")):
            flush()
            continue
        para_lines.append(stripped)
    flush()


def process_latex(document, text: str, image_dir: Path, citation_map: dict[str, str], references: list[str]):
    state = {
        "chapter": 0,
        "section": 0,
        "subsection": 0,
        "subsubsection": 0,
        "figure": 0,
        "table": 0,
        "equation": 0,
        "formula_image": 0,
        "pseudo_image": 0,
        "references": references,
    }
    for kind, block in parse_blocks(text):
        if kind == "text":
            process_text_chunk(document, block, state, citation_map)
        elif kind == "figure":
            add_figure(document, block, image_dir, state)
        elif kind == "table":
            add_table(document, block, citation_map, state)
        elif kind == "equation":
            add_equation(document, block, state)
        elif kind == "enumerate":
            add_enumerate(document, block, citation_map)
        elif kind == "researchitems":
            add_research_items(document, block, citation_map)
        elif kind == "center":
            add_center_block(document, block, state, citation_map)
        elif kind == "verbatim":
            screenshot = word_image_path(f"提示词{state['chapter']}-1.png")
            if screenshot:
                add_caption(document, f"提示词 {state['chapter']}-1 ", "终答提示词片段", before=True)
                add_screenshot(document, screenshot)
            else:
                add_body_paragraph(document, block, citation_map)


def configure_document(document: Document):
    styles = document.styles
    set_style_font(styles["Normal"], 12)
    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        set_style_font(styles[name], size, bold=True)
        clear_style_color(styles[name])
    for name, size in [("TOC 1", 12), ("TOC 2", 12), ("TOC 3", 12)]:
        style = get_or_create_paragraph_style(styles, name)
        set_style_font(style, size, italic=False, font_name=CHINESE_FONT)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
    if "List Number" in styles:
        set_style_font(styles["List Number"], 12)
    for style_name in ("CaptionNNU",):
        if style_name not in styles:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            set_style_font(style, 10.5)


def first_section_elements(body):
    elements = []
    for child in body:
        elements.append(child)
        if child.find(".//" + qn("w:sectPr")) is not None:
            break
    return elements


def relationship_content_type(rel_type: str) -> str | None:
    if rel_type.endswith("/header"):
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml"
    if rel_type.endswith("/footer"):
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml"
    if rel_type.endswith("/image"):
        return None
    return None


def add_content_type_override(content_types_root, part_name: str, content_type: str):
    existing = content_types_root.xpath(
        f'ct:Override[@PartName="{part_name}"]',
        namespaces={"ct": CT_NS},
    )
    if existing:
        return
    override = etree.Element(f"{{{CT_NS}}}Override")
    override.set("PartName", part_name)
    override.set("ContentType", content_type)
    content_types_root.append(override)


def add_content_type_default(content_types_root, extension: str, content_type: str):
    extension = extension.lower().lstrip(".")
    existing = content_types_root.xpath(
        f'ct:Default[@Extension="{extension}"]',
        namespaces={"ct": CT_NS},
    )
    if existing:
        return
    default = etree.Element(f"{{{CT_NS}}}Default")
    default.set("Extension", extension)
    default.set("ContentType", content_type)
    content_types_root.append(default)


def ensure_media_content_type(content_types_root, suffix: str):
    suffix = suffix.lower().lstrip(".")
    media_types = {
        "png": "image/png",
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "gif": "image/gif",
        "bmp": "image/bmp",
        "tif": "image/tiff",
        "tiff": "image/tiff",
    }
    if suffix in media_types:
        add_content_type_default(content_types_root, suffix, media_types[suffix])


def merge_source_styles(files: dict[str, bytes], source_files: dict[str, bytes]):
    if "word/styles.xml" not in files or "word/styles.xml" not in source_files:
        return
    target_root = etree_from_bytes(files["word/styles.xml"])
    source_root = etree_from_bytes(source_files["word/styles.xml"])
    target_ids = set(target_root.xpath("//w:style/@w:styleId", namespaces={"w": W_NS}))
    for style in source_root.xpath("//w:style", namespaces={"w": W_NS}):
        style_id = style.get(qn("w:styleId"))
        if style_id and style_id not in target_ids:
            target_root.append(deepcopy(style))
            target_ids.add(style_id)
    files["word/styles.xml"] = etree_to_bytes(target_root)


def merge_numbering_for_elements(files: dict[str, bytes], source_files: dict[str, bytes], elements):
    if "word/numbering.xml" not in source_files:
        return
    if "word/numbering.xml" not in files:
        files["word/numbering.xml"] = source_files["word/numbering.xml"]
        return

    source_num_ids = sorted({
        node.get(qn("w:val"))
        for el in elements
        for node in el.xpath(".//w:numPr/w:numId", namespaces={"w": W_NS})
        if node.get(qn("w:val"))
    })
    if not source_num_ids:
        return

    target_root = etree_from_bytes(files["word/numbering.xml"])
    source_root = etree_from_bytes(source_files["word/numbering.xml"])
    target_num_ids = {
        n.get(qn("w:numId"))
        for n in target_root.xpath("//w:num", namespaces={"w": W_NS})
        if n.get(qn("w:numId"))
    }
    target_abs_ids = {
        n.get(qn("w:abstractNumId"))
        for n in target_root.xpath("//w:abstractNum", namespaces={"w": W_NS})
        if n.get(qn("w:abstractNumId"))
    }
    next_num_id = max([int(x) for x in target_num_ids if x.isdigit()] or [0]) + 1
    next_abs_id = max([int(x) for x in target_abs_ids if x.isdigit()] or [0]) + 1

    source_nums = {
        n.get(qn("w:numId")): n
        for n in source_root.xpath("//w:num", namespaces={"w": W_NS})
        if n.get(qn("w:numId"))
    }
    source_abs = {
        n.get(qn("w:abstractNumId")): n
        for n in source_root.xpath("//w:abstractNum", namespaces={"w": W_NS})
        if n.get(qn("w:abstractNumId"))
    }
    num_id_map: dict[str, str] = {}
    abs_id_map: dict[str, str] = {}

    for old_num_id in source_num_ids:
        source_num = source_nums.get(old_num_id)
        if source_num is None:
            continue
        old_abs_ref = source_num.find(qn("w:abstractNumId"))
        old_abs_id = old_abs_ref.get(qn("w:val")) if old_abs_ref is not None else None
        if old_abs_id and old_abs_id not in abs_id_map:
            new_abs_id = str(next_abs_id)
            next_abs_id += 1
            abs_id_map[old_abs_id] = new_abs_id
            source_abstract = source_abs.get(old_abs_id)
            if source_abstract is not None:
                copied_abs = deepcopy(source_abstract)
                copied_abs.set(qn("w:abstractNumId"), new_abs_id)
                target_root.append(copied_abs)

        new_num_id = str(next_num_id)
        next_num_id += 1
        num_id_map[old_num_id] = new_num_id
        copied_num = deepcopy(source_num)
        copied_num.set(qn("w:numId"), new_num_id)
        copied_abs_ref = copied_num.find(qn("w:abstractNumId"))
        if copied_abs_ref is not None and old_abs_id in abs_id_map:
            copied_abs_ref.set(qn("w:val"), abs_id_map[old_abs_id])
        target_root.append(copied_num)

    for el in elements:
        for node in el.xpath(".//w:numPr/w:numId", namespaces={"w": W_NS}):
            old_num_id = node.get(qn("w:val"))
            if old_num_id in num_id_map:
                node.set(qn("w:val"), num_id_map[old_num_id])

    files["word/numbering.xml"] = etree_to_bytes(target_root)


def etree_from_bytes(data: bytes):
    return etree.fromstring(data)


def etree_to_bytes(root) -> bytes:
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def import_cover_page_from_docx(output_docx: Path, cover_docx: Path):
    """Replace the generated first section with the first section from cover_docx."""
    if not cover_docx.exists():
        return

    with ZipFile(output_docx, "r") as target_zip:
        files = {name: target_zip.read(name) for name in target_zip.namelist()}
    with ZipFile(cover_docx, "r") as source_zip:
        source_files = {name: source_zip.read(name) for name in source_zip.namelist()}

    target_doc = etree_from_bytes(files["word/document.xml"])
    source_doc = etree_from_bytes(source_files["word/document.xml"])
    target_body = target_doc.find(qn("w:body"))
    source_body = source_doc.find(qn("w:body"))
    target_cover = first_section_elements(target_body)
    source_cover = [deepcopy(el) for el in first_section_elements(source_body)]

    target_rels_path = "word/_rels/document.xml.rels"
    source_rels_path = "word/_rels/document.xml.rels"
    target_rels = etree_from_bytes(files[target_rels_path])
    source_rels = etree_from_bytes(source_files[source_rels_path])
    content_types = etree_from_bytes(files["[Content_Types].xml"])

    existing_ids = [rel.get("Id", "") for rel in target_rels]
    max_id = max([int(x[3:]) for x in existing_ids if x.startswith("rId") and x[3:].isdigit()] or [0])
    existing_names = set(files)
    rid_map: dict[str, str] = {}

    source_rel_by_id = {rel.get("Id"): rel for rel in source_rels}
    source_rids: list[str] = []
    for el in source_cover:
        source_rids.extend(el.xpath(".//@r:id", namespaces={"r": R_NS}))
    source_rids = sorted(set(source_rids))

    for old_rid in source_rids:
        rel = source_rel_by_id.get(old_rid)
        if rel is None:
            continue
        target = rel.get("Target")
        rel_type = rel.get("Type")
        mode = rel.get("TargetMode")
        max_id += 1
        new_rid = f"rId{max_id}"
        rid_map[old_rid] = new_rid

        new_rel = etree.Element(f"{{{REL_NS}}}Relationship")
        new_rel.set("Id", new_rid)
        new_rel.set("Type", rel_type)
        if mode:
            new_rel.set("TargetMode", mode)

        if mode == "External":
            new_rel.set("Target", target)
        else:
            source_part = str(Path("word") / target)
            suffix = Path(target).suffix
            stem = Path(target).stem
            new_target = f"cover_{stem}{suffix}"
            counter = 1
            while str(Path("word") / new_target) in existing_names:
                new_target = f"cover_{stem}_{counter}{suffix}"
                counter += 1
            target_part = str(Path("word") / new_target)
            if source_part in source_files:
                files[target_part] = source_files[source_part]
                existing_names.add(target_part)
                content_type = relationship_content_type(rel_type)
                if content_type:
                    add_content_type_override(content_types, "/" + target_part, content_type)
                if rel_type.endswith("/image"):
                    ensure_media_content_type(content_types, Path(target_part).suffix)
            new_rel.set("Target", new_target)
        target_rels.append(new_rel)

    for el in source_cover:
        for node in el.xpath(".//*[@r:id]", namespaces={"r": R_NS}):
            old_rid = node.get(qn("r:id"))
            if old_rid in rid_map:
                node.set(qn("r:id"), rid_map[old_rid])

    insert_at = list(target_body).index(target_cover[0])
    for el in target_cover:
        target_body.remove(el)
    for offset, el in enumerate(source_cover):
        target_body.insert(insert_at + offset, el)

    files["word/document.xml"] = etree_to_bytes(target_doc)
    files[target_rels_path] = etree_to_bytes(target_rels)
    files["[Content_Types].xml"] = etree_to_bytes(content_types)
    merge_source_styles(files, source_files)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = Path(tmp.name)
    try:
        with ZipFile(tmp_path, "w", ZIP_DEFLATED) as out_zip:
            for name, data in files.items():
                out_zip.writestr(name, data)
        shutil.move(str(tmp_path), output_docx)
    finally:
        if tmp_path.exists():
            tmp_path.unlink()


def remap_relationships_for_elements(
    elements,
    files: dict[str, bytes],
    source_files: dict[str, bytes],
    target_rels,
    source_rels,
    content_types,
):
    existing_ids = [rel.get("Id", "") for rel in target_rels]
    max_id = max([int(x[3:]) for x in existing_ids if x.startswith("rId") and x[3:].isdigit()] or [0])
    existing_names = set(files)
    rid_map: dict[str, str] = {}
    source_rel_by_id = {rel.get("Id"): rel for rel in source_rels}
    source_rids: list[str] = []
    for el in elements:
        source_rids.extend(el.xpath(".//@r:id", namespaces={"r": R_NS}))
        source_rids.extend(el.xpath(".//@r:embed", namespaces={"r": R_NS}))
        source_rids.extend(el.xpath(".//@r:link", namespaces={"r": R_NS}))

    for old_rid in sorted(set(source_rids)):
        rel = source_rel_by_id.get(old_rid)
        if rel is None:
            continue
        target = rel.get("Target")
        if not target or "NULL" in target:
            continue
        rel_type = rel.get("Type")
        mode = rel.get("TargetMode")
        max_id += 1
        new_rid = f"rId{max_id}"
        rid_map[old_rid] = new_rid

        new_rel = etree.Element(f"{{{REL_NS}}}Relationship")
        new_rel.set("Id", new_rid)
        new_rel.set("Type", rel_type)
        if mode:
            new_rel.set("TargetMode", mode)

        if mode == "External":
            new_rel.set("Target", target)
        else:
            source_part = str(Path("word") / target)
            target_path = Path(target)
            prefix = "content_"
            if target_path.parent != Path("."):
                new_target = str(target_path.parent / f"{prefix}{target_path.stem}{target_path.suffix}")
            else:
                new_target = f"{prefix}{target_path.stem}{target_path.suffix}"
            counter = 1
            while str(Path("word") / new_target) in existing_names:
                if target_path.parent != Path("."):
                    new_target = str(target_path.parent / f"{prefix}{target_path.stem}_{counter}{target_path.suffix}")
                else:
                    new_target = f"{prefix}{target_path.stem}_{counter}{target_path.suffix}"
                counter += 1

            target_part = str(Path("word") / new_target)
            if source_part in source_files:
                files[target_part] = source_files[source_part]
                existing_names.add(target_part)
                content_type = relationship_content_type(rel_type)
                if content_type:
                    add_content_type_override(content_types, "/" + target_part, content_type)
                if rel_type.endswith("/image"):
                    ensure_media_content_type(content_types, Path(target_part).suffix)
            new_rel.set("Target", new_target)
        target_rels.append(new_rel)

    for el in elements:
        for attr_name in ("r:id", "r:embed", "r:link"):
            for node in el.xpath(f".//*[@{attr_name}]", namespaces={"r": R_NS}):
                old_rid = node.get(qn(attr_name))
                if old_rid in rid_map:
                    node.set(qn(attr_name), rid_map[old_rid])


def remove_null_relationships(rels_root):
    for rel in list(rels_root):
        target = rel.get("Target") or ""
        if "NULL" in target:
            rel.getparent().remove(rel)


def build_output_from_template(template_docx: Path, generated_docx: Path, output_docx: Path):
    """Use template_docx as the package base and replace only content after its cover."""
    with ZipFile(template_docx, "r") as template_zip:
        files = {name: template_zip.read(name) for name in template_zip.namelist()}
    with ZipFile(generated_docx, "r") as generated_zip:
        generated_files = {name: generated_zip.read(name) for name in generated_zip.namelist()}

    target_doc = etree_from_bytes(files["word/document.xml"])
    generated_doc = etree_from_bytes(generated_files["word/document.xml"])
    target_body = target_doc.find(qn("w:body"))
    generated_body = generated_doc.find(qn("w:body"))

    target_cover = first_section_elements(target_body)
    generated_cover = first_section_elements(generated_body)
    generated_after_cover = [deepcopy(el) for el in list(generated_body)[len(generated_cover):]]

    target_rels_path = "word/_rels/document.xml.rels"
    generated_rels_path = "word/_rels/document.xml.rels"
    target_rels = etree_from_bytes(files[target_rels_path])
    generated_rels = etree_from_bytes(generated_files[generated_rels_path])
    content_types = etree_from_bytes(files["[Content_Types].xml"])

    remap_relationships_for_elements(
        generated_after_cover,
        files,
        generated_files,
        target_rels,
        generated_rels,
        content_types,
    )
    remove_null_relationships(target_rels)
    merge_source_styles(files, generated_files)
    merge_numbering_for_elements(files, generated_files, generated_after_cover)

    for child in list(target_body)[len(target_cover):]:
        target_body.remove(child)
    for el in generated_after_cover:
        target_body.append(el)

    files["word/document.xml"] = etree_to_bytes(target_doc)
    files[target_rels_path] = etree_to_bytes(target_rels)
    files["[Content_Types].xml"] = etree_to_bytes(content_types)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = Path(tmp.name)
    try:
        with ZipFile(tmp_path, "w", ZIP_DEFLATED) as out_zip:
            for name, data in files.items():
                out_zip.writestr(name, data)
        shutil.move(str(tmp_path), output_docx)
    finally:
        if tmp_path.exists():
            tmp_path.unlink()


def normalize_r_fonts(r_pr):
    if r_pr is None:
        return
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = etree.Element(qn("w:rFonts"))
        r_pr.insert(0, r_fonts)
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        r_fonts.attrib.pop(qn(f"w:{attr}"), None)
    r_fonts.set(qn("w:ascii"), WESTERN_FONT)
    r_fonts.set(qn("w:hAnsi"), WESTERN_FONT)
    r_fonts.set(qn("w:eastAsia"), CHINESE_FONT)
    r_fonts.set(qn("w:cs"), WESTERN_FONT)


def force_not_italic(r_pr):
    if r_pr is None:
        return
    for tag in ("w:i", "w:iCs"):
        node = r_pr.find(qn(tag))
        if node is None:
            node = etree.Element(qn(tag))
            r_pr.append(node)
        node.set(qn("w:val"), "0")


def sanitize_docx_fonts(docx_path: Path):
    """Remove template font leftovers so Word only sees SimSun and Times New Roman."""
    with ZipFile(docx_path, "r") as in_zip:
        files = {name: in_zip.read(name) for name in in_zip.namelist()}

    xml_parts = [
        name for name in files
        if name.startswith("word/")
        and name.endswith(".xml")
        and not name.startswith("word/media/")
        and name != "word/fontTable.xml"
    ]
    for name in xml_parts:
        try:
            root = etree_from_bytes(files[name])
        except Exception:
            continue
        changed = False
        for r_pr in root.xpath(".//w:rPr", namespaces={"w": W_NS}):
            normalize_r_fonts(r_pr)
            changed = True
        for style in root.xpath(".//w:style", namespaces={"w": W_NS}):
            style_id = style.get(qn("w:styleId")) or ""
            style_name = ""
            name_el = style.find(qn("w:name"))
            if name_el is not None:
                style_name = name_el.get(qn("w:val")) or ""
            if (
                style_id in {"TOC1", "TOC2", "TOC3", "TOCHeading"}
                or style_name in {"TOC 1", "TOC 2", "TOC 3", "TOC Heading", "Hyperlink"}
                or "Hyperlink" in style_id
            ):
                r_pr = style.find(qn("w:rPr"))
                if r_pr is None:
                    r_pr = etree.Element(qn("w:rPr"))
                    style.append(r_pr)
                normalize_r_fonts(r_pr)
                force_not_italic(r_pr)
                changed = True
        if changed:
            files[name] = etree_to_bytes(root)

    font_table = etree.Element(qn("w:fonts"), nsmap={"w": W_NS})
    for font_name, family in ((WESTERN_FONT, "roman"), (CHINESE_FONT, "auto")):
        font = etree.Element(qn("w:font"))
        font.set(qn("w:name"), font_name)
        charset = etree.Element(qn("w:charset"))
        charset.set(qn("w:val"), "86" if font_name == CHINESE_FONT else "00")
        family_el = etree.Element(qn("w:family"))
        family_el.set(qn("w:val"), family)
        pitch = etree.Element(qn("w:pitch"))
        pitch.set(qn("w:val"), "default" if font_name == CHINESE_FONT else "variable")
        font.extend([charset, family_el, pitch])
        font_table.append(font)
    files["word/fontTable.xml"] = etree_to_bytes(font_table)

    relationships = etree.Element(f"{{{REL_NS}}}Relationships")
    files["word/_rels/fontTable.xml.rels"] = etree_to_bytes(relationships)

    for name in list(files):
        if name.startswith("word/fonts/"):
            del files[name]

    if "[Content_Types].xml" in files:
        content_types = etree_from_bytes(files["[Content_Types].xml"])
        for el in list(content_types):
            part_name = el.get("PartName") or ""
            if part_name.startswith("/word/fonts/"):
                content_types.remove(el)
        files["[Content_Types].xml"] = etree_to_bytes(content_types)

    if "word/settings.xml" in files:
        settings = files["word/settings.xml"].decode("utf-8", errors="ignore")
        for old in ("Calibri", "Cambria Math", "Cambria"):
            settings = settings.replace(old, WESTERN_FONT)
        files["word/settings.xml"] = settings.encode("utf-8")

    if "word/theme/theme1.xml" in files:
        theme = files["word/theme/theme1.xml"].decode("utf-8", errors="ignore")
        theme = re.sub(r'typeface="[^"]*"', f'typeface="{WESTERN_FONT}"', theme)
        files["word/theme/theme1.xml"] = theme.encode("utf-8")

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = Path(tmp.name)
    try:
        with ZipFile(tmp_path, "w", ZIP_DEFLATED) as out_zip:
            for name, data in files.items():
                out_zip.writestr(name, data)
        shutil.move(str(tmp_path), docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink()


def main():
    global REF_MAP, ASSET_DIR, WORD_IMAGE_DIR
    parser = argparse.ArgumentParser()
    parser.add_argument("main_tex", type=Path)
    parser.add_argument("output_docx", type=Path)
    parser.add_argument("--cover-docx", type=Path, default=None)
    args = parser.parse_args()

    paper_dir = args.main_tex.parent
    ASSET_DIR = paper_dir / "_docx_generated_assets"
    WORD_IMAGE_DIR = paper_dir / "images_word"
    main_text = args.main_tex.read_text(encoding="utf-8")
    content_text = (paper_dir / "content.tex").read_text(encoding="utf-8")
    cn_text, cn_keywords, en_text, en_keywords, after_content = split_main_parts(main_text)
    full_body_text = content_text + "\n" + after_content
    REF_MAP = build_ref_map(full_body_text)
    citation_map = build_citation_map(paper_dir / "main.bbl")
    references = read_bibliography(paper_dir / "main.bbl")

    meta = {
        "title": get_macro(main_text, "thesistitle"),
        "author": get_macro(main_text, "thesisauthor"),
        "id": get_macro(main_text, "thesisid"),
        "major": get_macro(main_text, "thesismajor"),
        "college": get_macro(main_text, "thesiscollege"),
        "advisor": get_macro(main_text, "thesisadvisor"),
    }
    document = Document()
    configure_document(document)
    first_section = document.sections[0]
    set_header_footer(first_section, enabled=False)

    add_cover(document, meta)

    front = document.add_section(WD_SECTION_START.NEW_PAGE)
    set_header_footer(front, enabled=True, page_fmt="upperRoman", start=1)
    add_abstracts(document, cn_text, cn_keywords, en_text, en_keywords, citation_map)
    document.add_page_break()
    add_toc(document)

    body = document.add_section(WD_SECTION_START.CONTINUOUS)
    set_header_footer(body, enabled=True, page_fmt="decimal", start=1)
    process_latex(document, full_body_text, paper_dir / "images", citation_map, references)

    cover_docx = args.cover_docx or (paper_dir / "时子延毕业论文.docx")
    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        generated_path = Path(tmp.name)
    try:
        document.save(generated_path)
        if cover_docx.exists():
            build_output_from_template(cover_docx, generated_path, args.output_docx)
        else:
            shutil.move(str(generated_path), args.output_docx)
        sanitize_docx_fonts(args.output_docx)
    finally:
        if generated_path.exists():
            generated_path.unlink()
    print(f"Wrote {args.output_docx}")


if __name__ == "__main__":
    main()
